from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import RGBColor


class DocDeal:
    def deal_test(self):
        doc = Document()
        # 设置字体样式
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # ------添加文档标题-------
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(" 程序员有话说")
        font = run.font
        #  设置字体大小
        font.size = Pt(24)
        #  设置水平居中
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # ------添加一段话-------
        content = "这是一个最好的时代，也是一个最坏的时代。" \
                  "好的是众多程序员都加入通过文字表达自己想法的步伐，" \
                  "好的是依然围着技术转，始终不敢释放自我。你若不信，请听听他们的对话。"
        paragraph = doc.add_paragraph(content)
        paragraph_format = paragraph.paragraph_format
        #  第一行左边缩进
        paragraph_format.first_line_indent = Inches(0.3)
        # -----添加一个小标题------
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("同程序员聊天")
        paragraph_format = paragraph.paragraph_format
        #  段前
        paragraph_format.space_after = Pt(15)
        #  段后
        paragraph_format.space_before = Pt(2)
        #  字体加粗和下划线
        font = run.font
        font.bold = True
        # ----跟销售人员聊天-----
        content = "大家好，我是西门吹水，做销售的，现去研发部找程旭猿聊聊天。" \
                  "你好，程旭猿，在忙什么呢?\n研究Python技术当中。\n"
        paragraph = doc.add_paragraph(content)
        #  添加下划线
        run = paragraph.add_run("什么是派森来的?麻烦介绍一下。")
        font = run.font
        font.underline = True
        #  插入表格和内容
        table = doc.add_table(rows=3, cols=2, style="Medium Grid 1 Accent 1")
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "跨平台编程语言"
        table.cell(1, 0).text = "跨平台"
        table.cell(1, 1).text = "Windows、macOsS、Ubuntu等"
        table.cell(2, 0).text = "用途"
        table.cell(2, 1).text = "人工智能、Web、桌面系统..."
        # -----设置字体颜色------
        doc.add_paragraph("西门吹水：我一句都没听懂，怎么办法呀?")
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("程旭猿：给你一张图片，自己体会去。")
        font = run.font
        font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        # -----添加图片，设置图片大小------
        doc.add_picture(r"D:/bpz/python/actual/res/image/voilet.png", width=Inches(6.25))
        # ------保存word文档到当前目录下-------
        doc.save('D:/bpz/python/actual/download/doc/test.docx')


if __name__ == '__main__':
    DocDeal().deal_test()
